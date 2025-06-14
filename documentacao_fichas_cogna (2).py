# -*- coding: utf-8 -*-
"""DOCUMENTACAO_FICHAS_COGNA

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1s3wjwWNWAGPZbm9gIF_YNeBgTjs2IA0c
"""

!pip install pandas openpyxl python-docx  # Instalação (somente para Google Colab)

"""<small>

**Parte 1 — Importação das bibliotecas**

Importa todas as ferramentas que o código precisa:  
</small>
- `pandas` para trabalhar com planilhas.  
- `docx` para manipular documentos Word.  
- `os` para criar pastas e acessar arquivos.  
- `zipfile` para criar arquivos compactados (.zip).  
- `re` para modificar textos automaticamente.  
</small>


"""

# Importa as bibliotecas necessárias

import pandas as pd  # Trabalha com planilhas (Excel)
from docx import Document  # Manipula arquivos do Word (.docx)
import os  # Permite criar pastas e acessar arquivos
import zipfile  # Permite criar arquivos compactados (.zip)
import re  # Usada para modificar textos automaticamente

"""<small>

**Parte 2 — Definição dos caminhos**  
<small>
Define onde estão:  
- A planilha com os dados.  
- O modelo da ficha e da folha de rosto.  
- A pasta onde os arquivos preenchidos serão salvos.  
Isso garante que o programa saiba onde buscar e salvar arquivos.  
</small>

"""

# Define o caminho dos arquivos que serão usados

caminho_planilha = "/content/BASE_DE_FICHAS_PARA_AUTOMAÇÃO (3).xlsx"  # Arquivo Excel com os dados
caminho_template_ficha = "/content/Modelo_WBA0048_v3_Liderança_FC teste.docx"  # Modelo da ficha
caminho_template_modelo = "/content/Modelo_WBA0048_v3_Liderança (2).docx"  # Modelo da folha de rosto
pasta_saida = "fichas_final_com_tabelas"  # Pasta onde serão salvos os arquivos gerados

"""<small>

**Parte 3 — Criação da pasta de saída**  
<small>
Cria a pasta onde os arquivos gerados serão salvos.  
Se a pasta já existir, não dá erro.  
Garante que o programa possa rodar várias vezes sem problemas.  

</small>

"""

# Cria a pasta onde os arquivos serão salvos
# Se a pasta já existir, não faz nada

os.makedirs(pasta_saida, exist_ok=True)

"""<small>

**Parte 4 — Formatar o nome do autor**  
<small>
Transforma o nome no formato:  
- De: "Maria Silva"  
- Para: "Silva, Maria"  

Padrão obrigatório para fichas catalográficas.  
Se o nome estiver vazio ou for só uma palavra, mantém como está.  
</small>

"""

# Esta função pega o nome do autor e transforma no formato "Sobrenome, Nome"
# Exemplo: "Maria Silva" → "Silva, Maria"

def formatar_autor(nome):
    if pd.isna(nome) or not str(nome).strip():  # Se o nome estiver vazio
        return ""
    partes = nome.strip().split()  # Divide o nome em partes
    if len(partes) < 2:  # Se tiver só uma parte, retorna como está
        return nome.strip()
    sobrenome = partes[-1]  # Última parte é o sobrenome
    restante = " ".join(partes[:-1])  # O resto é o nome
    return f"{sobrenome}, {restante}"  # Junta no formato correto

"""<small>

**Parte 5 — Limpeza dos nomes de arquivos**  

Remove caracteres proibidos nos nomes de arquivos (`:`, `*`, `?`, etc.).  
Se o nome estiver vazio, usa "SemNome" para evitar falhas.  
</small>

"""

# Esta função remove caracteres proibidos em nomes de arquivos
# Exemplo: ":" ou "*" não podem estar no nome do arquivo

def limpar_nome_arquivo(texto):
    if pd.isna(texto) or str(texto).lower() == 'nan' or not str(texto).strip():  # Se estiver vazio
        return 'SemNome'
    texto = str(texto)
    texto = texto.replace(":", "")  # Remove dois-pontos
    texto = re.sub(r'[\\/*?"<>|]', "", texto)  # Remove outros caracteres proibidos
    return texto.strip() or 'SemNome'

"""<small>

**Parte 6 — Formatar palavras-chave**  

Deixa a primeira letra das palavras-chave maiúscula, por exemplo:  
- De: "ciência"  
- Para: "Ciência"  
Mantém a padronização e facilita a indexação.  
</small>

"""

# Deixa a primeira letra da palavra-chave maiúscula
# Exemplo: "ciência" → "Ciência"

def formatar_palavrachave(texto):
    if pd.isna(texto) or not str(texto).strip():  # Se estiver vazio
        return ""
    return str(texto).strip().capitalize()

"""<small>

**Parte 7 — Carregar a planilha de dados**  

Lê a planilha Excel com todos os dados dos documentos.  
Preenche valores vazios com strings vazias ('') para evitar erros.  
Cria a coluna 'AUTOR_FORMATADO' para uso posterior.  
</small>

"""

# Aqui o programa abre a planilha Excel com os dados
# E cria uma nova coluna com o nome do autor já formatado

try:
    df = pd.read_excel(caminho_planilha)  # Lê a planilha
    df.fillna('', inplace=True)  # Preenche valores vazios com ''
    df['AUTOR_FORMATADO'] = df['AUTOR1'].apply(formatar_autor)  # Formata o nome do autor
except Exception as e:
    print("❌ Erro ao carregar a planilha.")
    print(f"Detalhes: {e}")
    raise e  # Para o código se der erro

"""<small>

**Parte 8 — Substituição de textos no Word**  

Localiza marcadores como `<AUTOR>`, `<TITULO>`, etc., e substitui pelos dados reais.  
Se certos dados opcionais (como subtítulo ou volume) estiverem vazios, remove os trechos automaticamente.  
</small>

"""

# Substitui os textos marcados com <PLACEHOLDER> pelo conteúdo real
# Exemplo: <AUTOR> será trocado por "Silva, Maria"

def substituir_em_paragrafo(par, dados):
    if not par.text.strip():  # Se o parágrafo estiver vazio, não faz nada
        return
    novo_texto = par.text
    for chave, valor in dados.items():  # Para cada chave, troca pelo valor real
        novo_texto = novo_texto.replace(chave, str(valor) if pd.notna(valor) else "")
    if not dados.get('<SUBTITULO>'):  # Se não tiver subtítulo, remove o marcador
        novo_texto = re.sub(r':\s*<SUBTITULO>', '', novo_texto)
    if not dados.get('<VOLUME>'):  # Se não tiver volume, remove o marcador
        novo_texto = re.sub(r':\s*v\.\s*<VOLUME>', '', novo_texto)
    if novo_texto != par.text:
        for i in range(len(par.runs)):
            par.runs[i].text = ""  # Limpa o texto original
        par.text = novo_texto  # Coloca o novo texto

"""<small>

**Parte 9 — Substituição completa no documento**  

Aplica a substituição em todas as partes do documento Word:  
- Parágrafos.  
- Tabelas.  
Garante que todas as informações sejam preenchidas automaticamente.  
</small>


"""

# Aplica a substituição em todas as partes do documento
# Inclusive nas tabelas

def substituir_placeholders(doc, dados):
    for par in doc.paragraphs:
        substituir_em_paragrafo(par, dados)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    substituir_em_paragrafo(par, dados)

"""<small>

**Parte 10 — Criação do log de erros**  

Cria um arquivo `erros_log.txt` para registrar qualquer erro durante o processo.  
Útil para consultar posteriormente e verificar se algo falhou.  
</small>

"""

# Cria um arquivo para anotar qualquer erro que acontecer durante o processo

with open("erros_log.txt", "w") as log:
    log.write("LOG DE ERROS - GERAÇÃO DE FICHAS\n\n")

"""<small>

**Parte 11 — Processamento de cada linha da planilha**  
Para cada registro:  
- Formata dados.  
- Preenche modelos (ficha e folha de rosto).  
- Salva os arquivos em uma pasta específica.  
Se não houver autor, pula a linha.  
Se houver erro, registra no log e continua.  
</small>


"""

# Aqui o programa começa a criar as fichas para cada linha da planilha

for i, linha in df.iterrows():
    if not str(linha.get('AUTOR1', '')).strip():  # Se não tiver autor, pula
        continue

    try:
        # Formata o volume, se tiver
        volume_formatado = ""
        if str(linha.get('VOLUME', '')).strip():
            try:
                volume_formatado = f": v. {int(float(linha['VOLUME']))}"
            except:
                volume_formatado = ""

        # Ajusta título e subtítulo
        titulo_valor = str(linha.get('TITULO', '')).strip()
        subtitulo_valor = str(linha.get('SUBTITULO', '')).strip()
        if not subtitulo_valor and titulo_valor.endswith(":"):
            titulo_valor = titulo_valor[:-1].strip()
        subtitulo_formatado = f": {subtitulo_valor}" if subtitulo_valor else ""

        # Formata palavras-chave
        palavrachave1 = formatar_palavrachave(linha.get('PALAVRACHAVE1', ''))
        palavrachave2 = formatar_palavrachave(linha.get('PALAVRACHAVE2', ''))
        palavrachave3 = formatar_palavrachave(linha.get('PALAVRACHAVE3', ''))

        # Prepara os dados para substituir no Word
        substituicoes = {
            "<COORDENADOR>": linha.get('COORDENADOR', ''),
            "<REVISOR>": linha.get('REVISOR', ''),
            "<AUTOR>": linha['AUTOR_FORMATADO'],
            "<AUTOR1>": linha.get('AUTOR1', ''),
            "<AUTOR2>": linha.get('AUTOR2', ''),
            "<AUTOR3>": linha.get('AUTOR3', ''),
            "<CUTTER>": linha.get('CUTTER', ''),
            "<TITULO>": titulo_valor,
            "<SUBTITULO>": subtitulo_formatado,
            "<PAGINA>": str(int(float(linha['PAGINA']))) if str(linha.get('PAGINA', '')).strip() else '',
            "<ISBN>": linha.get('ISBN', ''),
            "<PALAVRACHAVE1>": palavrachave1,
            "<PALAVRACHAVE2>": palavrachave2,
            "<PALAVRACHAVE3>": palavrachave3,
            "<CDD>": linha.get('CDD', ''),
            "<VOLUME>": volume_formatado,
            "<NOME1>": linha['AUTOR_FORMATADO']
        }

        # Define o nome da pasta onde vai salvar
        autor_nome = limpar_nome_arquivo(linha['AUTOR_FORMATADO'])
        titulo_nome = limpar_nome_arquivo(titulo_valor)
        nome_pasta = f"{autor_nome} - {titulo_nome}"
        caminho_pasta = os.path.join(pasta_saida, nome_pasta)
        os.makedirs(caminho_pasta, exist_ok=True)

        # Gera a ficha
        doc_ficha = Document(caminho_template_ficha)
        substituir_placeholders(doc_ficha, substituicoes)
        doc_ficha.save(os.path.join(caminho_pasta, "Ficha Catalográfica.docx"))

        # Gera o modelo
        doc_modelo = Document(caminho_template_modelo)
        substituir_placeholders(doc_modelo, substituicoes)
        doc_modelo.save(os.path.join(caminho_pasta, "Modelo Preenchido.docx"))

        print(f"✅ Ficha gerada com sucesso: {nome_pasta}")

    except Exception as e:
        print(f"❌ Erro ao processar linha {i+1} - {linha.get('TITULO', 'Sem Título')}")
        print(f"    ➤ Detalhes: {e}")
        with open("erros_log.txt", "a") as log:
            log.write(f"Linha {i+1} - {linha.get('TITULO', 'Sem Título')}\n")
            log.write(f"Erro: {str(e)}\n\n")

"""<small>

**Parte 12 — Compactação dos arquivos gerados**  

Após gerar todas as fichas, junta tudo em um arquivo `.zip`.  
Facilita o envio e armazenamento dos arquivos criados.  
Se ocorrer erro, anota no log.  
</small>

"""

# Depois de criar todos os arquivos, junta tudo num ZIP

try:
    caminho_zip = "fichas_e_modelos.zip"
    with zipfile.ZipFile(caminho_zip, 'w') as zipf:
        for raiz, _, arquivos in os.walk(pasta_saida):
            for arquivo in arquivos:
                caminho_completo = os.path.join(raiz, arquivo)
                caminho_relativo = os.path.relpath(caminho_completo, pasta_saida)
                zipf.write(caminho_completo, caminho_relativo)
    print(f"📦 Arquivo compactado com sucesso: {caminho_zip}")
except Exception as e:
    print("❌ Erro ao criar o arquivo ZIP.")
    print(f"Detalhes: {e}")
    with open("erros_log.txt", "a") as log:
        log.write(f"Erro ao criar o ZIP: {str(e)}\n")

"""<small>

**Parte 13 — Finalização**

Exibe a mensagem "🚀 Processo finalizado" indicando que o processo terminou com sucesso.  
Tudo foi processado, salvo e compactado.  
</small>

"""

# Mostra mensagem que terminou tudo

print("🚀 Processo finalizado.")

"""### EXCLUIR ARQUIVOS GERADOS


---



"""

# prompt: excluir todos os arquivos gerados pelo processamento do meu codigo

import os
import shutil

# Define a pasta a ser excluída
pasta_saida = "fichas_final_com_tabelas"
caminho_zip = "fichas_e_modelos"

# Exclui a pasta e seu conteúdo
if os.path.exists(pasta_saida):
    shutil.rmtree(pasta_saida)
    print(f"Pasta '{pasta_saida}' e seu conteúdo foram excluídos com sucesso.")
else:
    print(f"Pasta '{pasta_saida}' não existe.")

# Exclui o arquivo zip
if os.path.exists(caminho_zip):
    os.remove(caminho_zip)
    print(f"Arquivo '{caminho_zip}' foi excluído com sucesso.")
else:
    print(f"Arquivo '{caminho_zip}' não existe.")